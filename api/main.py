# api/main.py
from __future__ import annotations

import io
import logging
import uuid

from fastapi import FastAPI, File, HTTPException, Request, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse

from pathlib import Path

from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse

from api.document_store import StoredDocument, get_document, save_document
from api.qa_service import answer_question_over_document
from api.retrieval_service import embed_chunks, retrieve_top_chunks
from api.schemas import (
    DocumentStatusResponse,
    DocumentUploadResponse,
    QARequest,
    QAResponse,
    SearchRequest,
    SearchResponse,
    SearchResultResponse,
)
from api.summarization_service import summarize_document
from ml.data.text_helpers import chunk_legal_text_with_offsets

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    from docx import Document
except ImportError:
    Document = None


logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI(title="Contract RAG API")

BASE_DIR = Path(__file__).resolve().parents[1]
FRONTEND_DIST_DIR = BASE_DIR / "frontend_dist"
ASSETS_DIR = FRONTEND_DIST_DIR / "assets"

app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:5173",
        "http://127.0.0.1:5173",
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

if ASSETS_DIR.exists():
    app.mount("/assets", StaticFiles(directory=ASSETS_DIR), name="assets")

@app.exception_handler(HTTPException)
async def http_exception_handler(request: Request, exc: HTTPException) -> JSONResponse:
    message = exc.detail if isinstance(exc.detail, str) else "Request failed."
    return JSONResponse(
        status_code=exc.status_code,
        content={"error": message},
    )


@app.exception_handler(Exception)
async def unhandled_exception_handler(request: Request, exc: Exception) -> JSONResponse:
    logger.exception("Unhandled server error: %s", exc)
    return JSONResponse(
        status_code=500,
        content={"error": "Internal server error."},
    )


@app.get("/health")
def health_check() -> dict[str, str]:
    return {
        "status": "ok",
        "message": "API is healthy.",
    }

@app.get("/{full_path:path}")
def serve_spa(full_path: str):
    if full_path.startswith("docs") or full_path.startswith("openapi") or full_path.startswith("redoc"):
        raise HTTPException(status_code=404, detail="Not found.")

    index_path = FRONTEND_DIST_DIR / "index.html"
    if index_path.exists():
        return FileResponse(index_path)

    raise HTTPException(status_code=404, detail="Frontend not found.")


def extract_text_from_txt(file_bytes: bytes) -> str:
    return file_bytes.decode("utf-8", errors="ignore").strip()


def extract_text_from_pdf(file_bytes: bytes) -> str:
    if fitz is None:
        raise HTTPException(
            status_code=500,
            detail="PyMuPDF is not installed.",
        )

    text_parts: list[str] = []
    pdf_stream = io.BytesIO(file_bytes)

    try:
        with fitz.open(stream=pdf_stream, filetype="pdf") as pdf:
            for page in pdf:
                text_parts.append(page.get_text())
    except Exception as exc:
        raise HTTPException(
            status_code=400,
            detail="Failed to read PDF file.",
        ) from exc

    return "\n".join(text_parts).strip()


def extract_text_from_docx(file_bytes: bytes) -> str:
    if Document is None:
        raise HTTPException(
            status_code=500,
            detail="python-docx is not installed.",
        )

    doc_stream = io.BytesIO(file_bytes)

    try:
        doc = Document(doc_stream)
    except Exception as exc:
        raise HTTPException(
            status_code=400,
            detail="Failed to read DOCX file.",
        ) from exc

    return "\n".join(paragraph.text for paragraph in doc.paragraphs).strip()


def extract_text(filename: str, file_bytes: bytes) -> str:
    lowered = filename.lower()

    if lowered.endswith(".txt"):
        return extract_text_from_txt(file_bytes)

    if lowered.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)

    if lowered.endswith(".docx"):
        return extract_text_from_docx(file_bytes)

    raise HTTPException(
        status_code=400,
        detail="Invalid file type. Please upload a .txt, .pdf, or .docx file.",
    )


@app.post("/documents", response_model=DocumentUploadResponse)
async def upload_document(file: UploadFile = File(...)) -> DocumentUploadResponse:
    if not file.filename:
        raise HTTPException(
            status_code=400,
            detail="Uploaded file is missing a filename.",
        )

    file_bytes = await file.read()
    if not file_bytes:
        raise HTTPException(
            status_code=400,
            detail="Uploaded file is empty.",
        )

    extracted_text = extract_text(file.filename, file_bytes)
    if not extracted_text:
        raise HTTPException(
            status_code=400,
            detail="Could not extract any text from the uploaded file.",
        )

    raw_chunks = chunk_legal_text_with_offsets(extracted_text)
    if not raw_chunks:
        raise HTTPException(
            status_code=400,
            detail="No usable chunks could be created from the uploaded document.",
        )

    document_id = f"doc_{uuid.uuid4().hex[:12]}"

    chunks: list[dict[str, object]] = []
    for i, chunk in enumerate(raw_chunks):
        chunks.append(
            {
                "chunk_id": f"{document_id}_chunk_{i}",
                "text": str(chunk["text"]),
                "start_char": int(chunk["start_char"]),
                "end_char": int(chunk["end_char"]),
            }
        )

    chunk_texts = [str(chunk["text"]) for chunk in chunks]
    embeddings = embed_chunks(chunk_texts)
    summary = summarize_document(extracted_text)

    stored_document = StoredDocument(
        document_id=document_id,
        filename=file.filename,
        status="ready",
        extracted_text=extracted_text,
        summary=summary,
        chunks=chunks,
        embeddings=embeddings,
    )

    save_document(stored_document)

    return DocumentUploadResponse(
        document_id=document_id,
        status="ready",
        filename=file.filename,
        text_length=len(extracted_text),
        chunk_count=len(chunks),
        summary=summary,
        extracted_text=extracted_text,
    )


@app.get("/documents/{document_id}", response_model=DocumentStatusResponse)
def get_document_status(document_id: str) -> DocumentStatusResponse:
    document = get_document(document_id)
    if document is None:
        raise HTTPException(
            status_code=404,
            detail="Document not found.",
        )

    return DocumentStatusResponse(
        document_id=document.document_id,
        status=document.status,
        filename=document.filename,
        text_length=len(document.extracted_text),
        chunk_count=len(document.chunks),
        summary=document.summary,
        extracted_text=document.extracted_text,
    )


@app.post("/search", response_model=SearchResponse)
def search_document(request: SearchRequest) -> SearchResponse:
    document = get_document(request.document_id)
    if document is None:
        raise HTTPException(
            status_code=404,
            detail="Document not found.",
        )

    results = retrieve_top_chunks(
        document=document,
        query=request.query,
        top_k=request.top_k,
        candidate_k=max(20, request.top_k),
    )

    return SearchResponse(
        document_id=request.document_id,
        query=request.query,
        results=[
            SearchResultResponse(
                chunk_id=row.chunk_id,
                score=row.rerank_score,
                text=row.text,
                start_char=row.start_char,
                end_char=row.end_char,
            )
            for row in results
        ],
    )


@app.post("/qa", response_model=QAResponse)
def qa_document(request: QARequest) -> QAResponse:
    document = get_document(request.document_id)
    if document is None:
        raise HTTPException(
            status_code=404,
            detail="Document not found.",
        )

    result = answer_question_over_document(
        document=document,
        question=request.question,
        top_k=request.top_k,
    )

    return QAResponse(**result)